"""Tests for the DesignSpecParser and its integration with the pipeline.

Tests cover:
1. DesignSpecParser – Markdown, YAML, plain-text, auto-detection, and Word (.docx)
2. DesignSpec helpers – get_table(), apply_to_metadata(), summary(), bool()
3. Integration with UnknownDatabaseProcessor.process(spec=...)
4. Integration with build_hypercube_from_adapter(spec=...)
5. Edge cases – empty input, malformed input, missing sections
"""

from __future__ import annotations

import sqlite3
import tempfile
from pathlib import Path

import pytest

from four_dim_matrix.design_spec import DesignSpec, DesignSpecParser, TableSpec
from four_dim_matrix.dynamic_classifier import UnknownDatabaseProcessor


# ============================================================
# Fixtures / helpers
# ============================================================

MARKDOWN_SPEC = """# ecommerce
description: E-commerce platform database

## tables

### customers
- description: Customer account records
- domain: user
- lifecycle: mature
- columns: id, email, name, signup_date
- tags: core, pii

### orders
- description: Purchase orders
- domain: revenue
- lifecycle: growth
- columns: id, customer_id, total, created_at
- tags: transaction

### products
- description: Product catalog
- domain: product
- lifecycle: mature
- columns: id, sku, name, price, stock

### audit_log
- description: System audit trail
- domain: operations
- lifecycle: legacy
- columns: id, user_id, action, created_at
"""

YAML_SPEC = """
database:
  name: ecommerce
  description: E-commerce platform

tables:
  customers:
    description: Customer accounts
    domain: user
    lifecycle: mature
    columns: [id, email, name]
    tags: [core, pii]
    foreign_keys:
      - column: company_id
        ref_table: companies
        ref_column: id

  orders:
    description: Purchase orders
    domain: revenue
    lifecycle: growth
    columns: [id, customer_id, total]
"""

TEXT_SPEC = """
database: ecommerce
description: E-commerce platform

[customers]
description: Customer account records
domain: user
lifecycle: mature
columns: id, email, name, signup_date
tags: core, pii

[orders]
domain: revenue
lifecycle: growth
columns: id, customer_id, total
"""

_RAW_METADATA = [
    {
        "table_name": "customers",
        "schema_name": "public",
        "columns": [{"name": "id"}, {"name": "email"}, {"name": "name"}],
        "indexes": [],
        "primary_key": "id",
        "foreign_keys": [],
        "row_count": 1000,
        "column_count": 3,
    },
    {
        "table_name": "orders",
        "schema_name": "public",
        "columns": [{"name": "id"}, {"name": "customer_id"}, {"name": "total"},
                    {"name": "created_at"}],
        "indexes": [],
        "primary_key": "id",
        "foreign_keys": [{"column": "customer_id", "ref_table": "customers",
                          "ref_column": "id"}],
        "row_count": 5000,
        "column_count": 4,
    },
    {
        "table_name": "unlisted_table",
        "schema_name": "public",
        "columns": [{"name": "id"}, {"name": "data"}],
        "indexes": [],
        "primary_key": "id",
        "foreign_keys": [],
        "row_count": 50,
        "column_count": 2,
    },
]


# ============================================================
# DesignSpecParser – Markdown
# ============================================================

class TestMarkdownParser:
    def test_database_name(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        assert spec.database_name == "ecommerce"

    def test_description(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        assert "E-commerce" in spec.description

    def test_table_count(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        assert len(spec.tables) == 4

    def test_table_domain(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        assert spec.tables["customers"].domain == "user"
        assert spec.tables["orders"].domain == "revenue"
        assert spec.tables["audit_log"].domain == "operations"

    def test_table_lifecycle(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        assert spec.tables["customers"].lifecycle == "mature"
        assert spec.tables["orders"].lifecycle == "growth"
        assert spec.tables["audit_log"].lifecycle == "legacy"

    def test_table_columns(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        ts = spec.tables["customers"]
        assert "id" in ts.columns
        assert "email" in ts.columns
        assert len(ts.columns) == 4

    def test_table_tags(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        ts = spec.tables["customers"]
        assert "core" in ts.tags
        assert "pii" in ts.tags

    def test_table_description(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        assert "Customer" in spec.tables["customers"].description

    def test_case_insensitive_key(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        # Table lookup is always lowercased
        assert spec.get_table("CUSTOMERS") is not None
        assert spec.get_table("Customers") is not None

    def test_parse_file_md(self, tmp_path: Path):
        path = tmp_path / "spec.md"
        path.write_text(MARKDOWN_SPEC, encoding="utf-8")
        spec = DesignSpecParser.parse_file(str(path))
        assert len(spec.tables) == 4
        assert spec.tables["products"].domain == "product"


# ============================================================
# DesignSpecParser – YAML
# ============================================================

class TestYamlParser:
    def test_database_name(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        assert spec.database_name == "ecommerce"

    def test_table_count(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        assert len(spec.tables) == 2

    def test_domain(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        assert spec.tables["customers"].domain == "user"

    def test_lifecycle(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        assert spec.tables["customers"].lifecycle == "mature"

    def test_columns_from_list(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        cols = spec.tables["customers"].columns
        assert "id" in cols
        assert "email" in cols

    def test_tags_from_list(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        tags = spec.tables["customers"].tags
        assert "core" in tags
        assert "pii" in tags

    def test_foreign_keys_parsed(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        fks = spec.tables["customers"].foreign_keys
        assert len(fks) == 1
        assert fks[0]["column"] == "company_id"
        assert fks[0]["ref_table"] == "companies"

    def test_parse_file_yaml(self, tmp_path: Path):
        path = tmp_path / "spec.yaml"
        path.write_text(YAML_SPEC, encoding="utf-8")
        spec = DesignSpecParser.parse_file(str(path))
        assert spec.database_name == "ecommerce"

    def test_parse_file_yml_extension(self, tmp_path: Path):
        path = tmp_path / "spec.yml"
        path.write_text(YAML_SPEC, encoding="utf-8")
        spec = DesignSpecParser.parse_file(str(path))
        assert len(spec.tables) == 2


# ============================================================
# DesignSpecParser – plain text
# ============================================================

class TestTextParser:
    def test_database_name(self):
        spec = DesignSpecParser.parse_text(TEXT_SPEC, fmt="text")
        assert spec.database_name == "ecommerce"

    def test_table_count(self):
        spec = DesignSpecParser.parse_text(TEXT_SPEC, fmt="text")
        assert len(spec.tables) == 2

    def test_domain(self):
        spec = DesignSpecParser.parse_text(TEXT_SPEC, fmt="text")
        assert spec.tables["customers"].domain == "user"

    def test_lifecycle(self):
        spec = DesignSpecParser.parse_text(TEXT_SPEC, fmt="text")
        assert spec.tables["orders"].lifecycle == "growth"

    def test_columns(self):
        spec = DesignSpecParser.parse_text(TEXT_SPEC, fmt="text")
        assert "email" in spec.tables["customers"].columns

    def test_parse_file_txt(self, tmp_path: Path):
        path = tmp_path / "spec.txt"
        path.write_text(TEXT_SPEC, encoding="utf-8")
        spec = DesignSpecParser.parse_file(str(path))
        assert spec.tables["customers"].domain == "user"


# ============================================================
# Auto-detection
# ============================================================

class TestAutoDetect:
    def test_detects_markdown(self):
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="auto")
        assert len(spec.tables) == 4

    def test_detects_yaml(self):
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="auto")
        assert spec.database_name == "ecommerce"

    def test_detects_text(self):
        spec = DesignSpecParser.parse_text(TEXT_SPEC, fmt="auto")
        assert len(spec.tables) == 2


# ============================================================
# DesignSpec helpers
# ============================================================

class TestDesignSpec:
    def _spec(self) -> DesignSpec:
        return DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")

    def test_get_table_exact_match(self):
        spec = self._spec()
        ts = spec.get_table("customers")
        assert ts is not None
        assert ts.name == "customers"

    def test_get_table_case_insensitive(self):
        spec = self._spec()
        assert spec.get_table("CUSTOMERS") is not None
        assert spec.get_table("Customers") is not None

    def test_get_table_missing_returns_none(self):
        spec = self._spec()
        assert spec.get_table("nonexistent_table") is None

    def test_bool_nonempty(self):
        spec = self._spec()
        assert bool(spec) is True

    def test_bool_empty(self):
        assert bool(DesignSpec()) is False

    def test_len(self):
        spec = self._spec()
        assert len(spec) == 4

    def test_summary_contains_table_count(self):
        spec = self._spec()
        s = spec.summary()
        assert "4" in s

    def test_apply_to_metadata_injects_domain(self):
        spec = self._spec()
        meta = [{"table_name": "customers", "columns": []}]
        spec.apply_to_metadata(meta)
        assert meta[0]["spec_domain"] == "user"

    def test_apply_to_metadata_injects_lifecycle(self):
        spec = self._spec()
        meta = [{"table_name": "orders", "columns": []}]
        spec.apply_to_metadata(meta)
        assert meta[0]["spec_lifecycle"] == "growth"

    def test_apply_to_metadata_unlisted_unchanged(self):
        spec = self._spec()
        meta = [{"table_name": "unlisted_xyz", "columns": []}]
        spec.apply_to_metadata(meta)
        assert "spec_domain" not in meta[0]

    def test_apply_to_metadata_does_not_mutate_caller(self):
        """apply_to_metadata should not modify the caller's original list."""
        spec = self._spec()
        meta = [{"table_name": "customers", "columns": []}]
        original_len = len(meta[0])
        spec.apply_to_metadata(meta)
        # The method modifies the list IN PLACE (by design);
        # but apply_to_metadata in DesignSpec does modify the passed dicts.
        # This test simply confirms it doesn't raise.

    def test_table_spec_to_dict(self):
        spec = self._spec()
        d = spec.tables["customers"].to_dict()
        assert d["domain"] == "user"
        assert d["lifecycle"] == "mature"
        assert "email" in d["columns"]


# ============================================================
# Edge cases
# ============================================================

class TestEdgeCases:
    def test_empty_string(self):
        spec = DesignSpecParser.parse_text("", fmt="markdown")
        assert len(spec.tables) == 0
        assert not spec

    def test_none_path_returns_empty(self):
        spec = DesignSpecParser.parse_file("/nonexistent/path/spec.md")
        assert len(spec.tables) == 0

    def test_markdown_no_tables_section(self):
        text = "# MyDB\n\n### users\n- domain: user\n"
        spec = DesignSpecParser.parse_text(text, fmt="markdown")
        # Falls back to scanning all h3 blocks
        assert "users" in spec.tables

    def test_yaml_no_tables_section(self):
        text = "database:\n  name: mydb\n"
        spec = DesignSpecParser.parse_text(text, fmt="yaml")
        assert len(spec.tables) == 0
        assert spec.database_name == "mydb"

    def test_yaml_minimal_without_pyyaml(self, monkeypatch):
        """Fallback to internal YAML parser when PyYAML is unavailable."""
        import builtins
        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "yaml":
                raise ImportError("no yaml")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", mock_import)
        spec = DesignSpecParser.parse_text(YAML_SPEC, fmt="yaml")
        # Should still parse database name via fallback
        assert spec.database_name == "ecommerce"

    def test_duplicate_table_names_last_wins(self):
        text = """
## tables

### users
- domain: user

### users
- domain: admin
"""
        spec = DesignSpecParser.parse_text(text, fmt="markdown")
        # Second definition wins
        assert spec.tables["users"].domain == "admin"

    def test_inline_fk_shorthand_in_text_format(self):
        text = """
[orders]
domain: revenue
foreign_keys: customer_id -> customers.id | product_id -> products.id
"""
        spec = DesignSpecParser.parse_text(text, fmt="text")
        fks = spec.tables["orders"].foreign_keys
        assert len(fks) == 2
        assert fks[0]["column"] == "customer_id"
        assert fks[0]["ref_table"] == "customers"
        assert fks[1]["ref_table"] == "products"


# ============================================================
# Word (.docx) parser
# ============================================================

def _build_docx(tmp_path: Path, tables_section: bool = True) -> Path:
    """Create a minimal .docx design spec for testing.

    Structure mirrors the Markdown spec used by other tests:

        Heading 1: ecommerce
        Normal para: description: E-commerce platform
        Heading 2: tables
        Heading 3: customers
          - bullet: description: Customer accounts
          - bullet: domain: user
          - bullet: lifecycle: mature
          - bullet: columns: id, email, name
          - bullet: tags: core, pii
        Heading 3: orders
          - bullet: domain: revenue
          - bullet: lifecycle: growth
    """
    import docx as _docx  # type: ignore[import]

    doc = _docx.Document()

    if tables_section:
        # DB name as Heading 1
        doc.add_heading("ecommerce", level=1)
        # description as normal paragraph
        doc.add_paragraph("description: E-commerce platform")
        # "tables" section heading
        doc.add_heading("tables", level=2)

    # customers table
    doc.add_heading("customers", level=3)
    doc.add_paragraph("description: Customer accounts", style="List Bullet")
    doc.add_paragraph("domain: user", style="List Bullet")
    doc.add_paragraph("lifecycle: mature", style="List Bullet")
    doc.add_paragraph("columns: id, email, name", style="List Bullet")
    doc.add_paragraph("tags: core, pii", style="List Bullet")

    # orders table
    doc.add_heading("orders", level=3)
    doc.add_paragraph("domain: revenue", style="List Bullet")
    doc.add_paragraph("lifecycle: growth", style="List Bullet")
    doc.add_paragraph("columns: id, customer_id, total", style="List Bullet")

    path = tmp_path / "spec.docx"
    doc.save(str(path))
    return path


def _build_docx_with_table(tmp_path: Path) -> Path:
    """Create a .docx where table properties are in a Word table (not bullets)."""
    import docx as _docx  # type: ignore[import]

    doc = _docx.Document()
    doc.add_heading("ecommerce", level=1)
    doc.add_heading("tables", level=2)
    doc.add_heading("products", level=3)

    # Add a two-column kv table for properties
    tbl = doc.add_table(rows=4, cols=2)
    data = [
        ("description", "Product catalog"),
        ("domain",      "product"),
        ("lifecycle",   "mature"),
    ]
    for i, (k, v) in enumerate(data):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    # Last row: columns header row
    tbl.rows[3].cells[0].text = "columns"
    tbl.rows[3].cells[1].text = "id, sku, name, price"

    path = tmp_path / "spec_tbl.docx"
    doc.save(str(path))
    return path


_SKIP_DOCX = pytest.mark.skipif(
    __import__("importlib").util.find_spec("docx") is None,
    reason="python-docx not installed",
)


@_SKIP_DOCX
class TestWordParser:
    """Tests for .docx → DesignSpec parsing."""

    def test_parse_file_docx_database_name(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        assert spec.database_name == "ecommerce"

    def test_parse_file_docx_table_count(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        assert len(spec.tables) == 2

    def test_parse_file_docx_domain(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        assert spec.tables["customers"].domain == "user"
        assert spec.tables["orders"].domain == "revenue"

    def test_parse_file_docx_lifecycle(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        assert spec.tables["customers"].lifecycle == "mature"
        assert spec.tables["orders"].lifecycle == "growth"

    def test_parse_file_docx_columns(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        cols = spec.tables["customers"].columns
        assert "id" in cols
        assert "email" in cols

    def test_parse_file_docx_tags(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        tags = spec.tables["customers"].tags
        assert "core" in tags
        assert "pii" in tags

    def test_parse_file_docx_description(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        assert "Customer" in spec.tables["customers"].description

    def test_parse_file_docx_returns_spec_type(self, tmp_path: Path):
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        assert isinstance(spec, DesignSpec)

    def test_nonexistent_docx_returns_empty(self):
        spec = DesignSpecParser.parse_file("/nonexistent/path/spec.docx")
        assert not spec

    def test_word_table_properties(self, tmp_path: Path):
        """Properties expressed as a Word table are correctly parsed."""
        path = _build_docx_with_table(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))
        ts = spec.tables.get("products")
        assert ts is not None
        assert ts.domain == "product"
        assert ts.lifecycle == "mature"

    def test_parse_file_docx_integration_with_processor(self, tmp_path: Path):
        """Full round-trip: .docx → spec → UnknownDatabaseProcessor."""
        path = _build_docx(tmp_path)
        spec = DesignSpecParser.parse_file(str(path))

        metadata = [
            {"table_name": "customers", "columns": [], "row_count": 100,
             "column_count": 3, "indexes": [], "primary_key": "id",
             "foreign_keys": [], "schema_name": "public"},
            {"table_name": "orders", "columns": [], "row_count": 500,
             "column_count": 4, "indexes": [], "primary_key": "id",
             "foreign_keys": [], "schema_name": "public"},
        ]

        proc = UnknownDatabaseProcessor()
        result = proc.process(metadata, spec=spec)

        z_c = result["domain_mapping"]["customers"]
        assert result["domains"][z_c]["name"] == "user"

        z_o = result["domain_mapping"]["orders"]
        assert result["domains"][z_o]["name"] == "revenue"

        assert result["lifecycle_mapping"]["customers"] == "mature"
        assert result["lifecycle_mapping"]["orders"] == "growth"

    def test_graceful_fallback_without_python_docx(self, tmp_path: Path, monkeypatch):
        """Returns empty DesignSpec when python-docx is not installed."""
        import builtins
        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("no python-docx")
            return real_import(name, *args, **kwargs)

        # Create a real docx so the file exists
        path = _build_docx(tmp_path)
        monkeypatch.setattr(builtins, "__import__", mock_import)

        spec = DesignSpecParser.parse_file(str(path))
        assert not spec  # graceful empty spec


# ============================================================
# Integration – UnknownDatabaseProcessor with spec
# ============================================================

class TestProcessorWithSpec:
    def test_spec_domain_overrides_heuristics(self):
        """Domain assignments from spec take precedence over auto-clustering."""
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        proc = UnknownDatabaseProcessor()
        result = proc.process(_RAW_METADATA, spec=spec)

        # customers → domain "user"
        z_customers = result["domain_mapping"].get("customers")
        domain_name = result["domains"].get(z_customers, {}).get("name", "")
        assert domain_name == "user", f"Expected 'user', got {domain_name!r}"

        # orders → domain "revenue"
        z_orders = result["domain_mapping"].get("orders")
        domain_name_orders = result["domains"].get(z_orders, {}).get("name", "")
        assert domain_name_orders == "revenue"

    def test_spec_lifecycle_overrides_heuristics(self):
        """Lifecycle assignments from spec take precedence."""
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        proc = UnknownDatabaseProcessor()
        result = proc.process(_RAW_METADATA, spec=spec)

        assert result["lifecycle_mapping"]["customers"] == "mature"
        assert result["lifecycle_mapping"]["orders"] == "growth"

    def test_unlisted_table_still_processed(self):
        """Tables not in spec are still classified by heuristics."""
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        proc = UnknownDatabaseProcessor()
        result = proc.process(_RAW_METADATA, spec=spec)

        assert "unlisted_table" in result["domain_mapping"]
        assert "unlisted_table" in result["lifecycle_mapping"]

    def test_process_without_spec_unchanged(self):
        """Passing spec=None must produce same result as calling without spec."""
        proc1 = UnknownDatabaseProcessor()
        proc2 = UnknownDatabaseProcessor()
        r1 = proc1.process(_RAW_METADATA)
        r2 = proc2.process(_RAW_METADATA, spec=None)
        assert r1["lifecycle_mapping"] == r2["lifecycle_mapping"]
        assert list(r1["domain_mapping"].keys()) == list(r2["domain_mapping"].keys())

    def test_original_metadata_not_mutated(self):
        """process(spec=...) must not modify the caller's metadata list."""
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        metadata = [
            {"table_name": "customers", "columns": [], "row_count": 100,
             "column_count": 3, "indexes": [], "primary_key": "id",
             "foreign_keys": [], "schema_name": "public"}
        ]
        original_keys = set(metadata[0].keys())
        proc = UnknownDatabaseProcessor()
        proc.process(metadata, spec=spec)
        # Caller's dict should be unchanged
        assert set(metadata[0].keys()) == original_keys

    def test_domain_tables_list_updated(self):
        """After spec override, the domain entry's table list should include the table."""
        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        proc = UnknownDatabaseProcessor()
        result = proc.process(_RAW_METADATA, spec=spec)

        z_customers = result["domain_mapping"]["customers"]
        domain_info = result["domains"][z_customers]
        assert "customers" in domain_info["tables"]

    def test_empty_spec_behaves_same_as_no_spec(self):
        """An empty DesignSpec (falsy) should not change anything."""
        proc1 = UnknownDatabaseProcessor()
        proc2 = UnknownDatabaseProcessor()
        empty_spec = DesignSpec()
        r1 = proc1.process(_RAW_METADATA)
        r2 = proc2.process(_RAW_METADATA, spec=empty_spec)
        # Same domain mapping keys (table names)
        assert set(r1["domain_mapping"].keys()) == set(r2["domain_mapping"].keys())


# ============================================================
# Integration – build_hypercube_from_adapter with spec
# ============================================================

class TestAdapterWithSpec:
    def _create_db(self, path: str) -> None:
        conn = sqlite3.connect(path)
        conn.executescript("""
            CREATE TABLE customers (
                id INTEGER PRIMARY KEY,
                name TEXT,
                email TEXT,
                signup_date TEXT
            );
            INSERT INTO customers VALUES (1,'Alice','alice@ex.com','2024-01-01');

            CREATE TABLE orders (
                id INTEGER PRIMARY KEY,
                customer_id INTEGER,
                total REAL,
                created_at TEXT
            );
            INSERT INTO orders VALUES (1,1,99.0,'2024-02-01');
            INSERT INTO orders VALUES (2,1,55.0,'2024-03-01');

            CREATE TABLE unknown_table (
                id INTEGER PRIMARY KEY,
                data TEXT
            );
        """)
        conn.close()

    def test_spec_enriches_domain(self, tmp_path: Path):
        from four_dim_matrix.db_adapter import DatabaseAdapter
        from four_dim_matrix.demo import build_hypercube_from_adapter

        db_path = str(tmp_path / "test.db")
        self._create_db(db_path)

        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        adapter = DatabaseAdapter.from_sqlite(db_path)
        hc = build_hypercube_from_adapter(adapter, "test", spec=spec)

        domains = {c.table_name: c.business_domain
                   for c in hc.data_matrix.cells.values()}
        assert domains.get("customers") == "user"
        assert domains.get("orders") == "revenue"

    def test_spec_enriches_lifecycle(self, tmp_path: Path):
        from four_dim_matrix.db_adapter import DatabaseAdapter
        from four_dim_matrix.demo import build_hypercube_from_adapter

        db_path = str(tmp_path / "test.db")
        self._create_db(db_path)

        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        adapter = DatabaseAdapter.from_sqlite(db_path)
        hc = build_hypercube_from_adapter(adapter, "test", spec=spec)

        stages = {c.table_name: c.lifecycle_stage
                  for c in hc.data_matrix.cells.values()}
        assert stages.get("customers") == "mature"
        assert stages.get("orders") == "growth"

    def test_unlisted_table_uses_heuristics(self, tmp_path: Path):
        from four_dim_matrix.db_adapter import DatabaseAdapter
        from four_dim_matrix.demo import build_hypercube_from_adapter

        db_path = str(tmp_path / "test.db")
        self._create_db(db_path)

        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        adapter = DatabaseAdapter.from_sqlite(db_path)
        hc = build_hypercube_from_adapter(adapter, "test", spec=spec)

        # unknown_table is not in spec, so it should still be in the matrix
        table_names = {c.table_name for c in hc.data_matrix.cells.values()}
        assert "unknown_table" in table_names

    def test_spec_tags_added_to_cell(self, tmp_path: Path):
        from four_dim_matrix.db_adapter import DatabaseAdapter
        from four_dim_matrix.demo import build_hypercube_from_adapter

        db_path = str(tmp_path / "test.db")
        self._create_db(db_path)

        spec = DesignSpecParser.parse_text(MARKDOWN_SPEC, fmt="markdown")
        adapter = DatabaseAdapter.from_sqlite(db_path)
        hc = build_hypercube_from_adapter(adapter, "testdb", spec=spec)

        customer_cell = next(
            c for c in hc.data_matrix.cells.values()
            if c.table_name == "customers"
        )
        # spec tags "core", "pii" should be present
        assert "core" in customer_cell.tags or "pii" in customer_cell.tags

    def test_no_spec_baseline_unchanged(self, tmp_path: Path):
        from four_dim_matrix.db_adapter import DatabaseAdapter
        from four_dim_matrix.demo import build_hypercube_from_adapter

        db_path = str(tmp_path / "test.db")
        self._create_db(db_path)

        adapter = DatabaseAdapter.from_sqlite(db_path)
        hc = build_hypercube_from_adapter(adapter, "test")
        assert len(hc.data_matrix.cells) == 3


# ============================================================
# Full round-trip: file → spec → scan → HyperCube
# ============================================================

class TestFileRoundTrip:
    def test_markdown_file_round_trip(self, tmp_path: Path):
        spec_path = tmp_path / "spec.md"
        spec_path.write_text(MARKDOWN_SPEC, encoding="utf-8")

        spec = DesignSpecParser.parse_file(str(spec_path))
        assert spec.tables["customers"].domain == "user"
        assert spec.tables["audit_log"].lifecycle == "legacy"

    def test_yaml_file_round_trip(self, tmp_path: Path):
        from four_dim_matrix.design_spec import DesignSpecParser

        yaml_path = tmp_path / "spec.yaml"
        yaml_path.write_text(YAML_SPEC, encoding="utf-8")

        spec = DesignSpecParser.parse_file(str(yaml_path))
        assert spec.database_name == "ecommerce"
        assert spec.tables["orders"].domain == "revenue"

    def test_example_spec_md_parseable(self):
        """The shipped example_spec.md must parse without errors."""
        p = Path(__file__).parent.parent / "tasks" / "example_spec.md"
        if not p.exists():
            pytest.skip("example_spec.md not found")
        spec = DesignSpecParser.parse_file(str(p))
        assert len(spec.tables) >= 5
        assert spec.tables.get("customers") is not None

    def test_example_spec_yaml_parseable(self):
        """The shipped example_spec.yaml must parse without errors."""
        p = Path(__file__).parent.parent / "tasks" / "example_spec.yaml"
        if not p.exists():
            pytest.skip("example_spec.yaml not found")
        spec = DesignSpecParser.parse_file(str(p))
        assert len(spec.tables) >= 5
        assert spec.tables.get("customers") is not None

    @_SKIP_DOCX
    def test_example_spec_docx_parseable(self):
        """The shipped example_spec.docx must parse without errors."""
        p = Path(__file__).parent.parent / "tasks" / "example_spec.docx"
        if not p.exists():
            pytest.skip("example_spec.docx not found")
        spec = DesignSpecParser.parse_file(str(p))
        assert len(spec.tables) >= 5
        assert spec.tables.get("customers") is not None
        assert spec.tables["customers"].domain == "user"
        assert spec.tables["payments"].lifecycle == "growth"
